#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Парсер DOCX документов
"""

import time
from pathlib import Path
from typing import Dict, Any, Optional
import logging

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base_parser import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    Парсер DOCX документов
    
    Поддерживает:
    - Извлечение текста с сохранением структуры
    - Извлечение таблиц
    - Метаданные документа
    """
    
    SUPPORTED_EXTENSIONS = [".docx"]
    
    def supports(self, file_path: Path) -> bool:
        """Проверка поддержки формата"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: Path) -> ParseResult:
        """
        Парсинг DOCX документа
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            Результат парсинга
        """
        start_time = time.time()
        result = ParseResult()
        
        try:
            self.validate_file(file_path)
            
            # Базовые метаданные
            result.metadata = self.extract_metadata(file_path)
            
            # Открытие документа
            doc = docx.Document(str(file_path))
            
            # Извлечение метаданных DOCX
            docx_metadata = self._extract_docx_metadata(doc)
            result.metadata.update(docx_metadata)
            
            # Извлечение текста
            text_parts = []
            tables = []
            
            for element in doc.element.body:
                # Параграфы
                if element.tag.endswith('p'):
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Таблицы
                elif element.tag.endswith('tbl'):
                    table = Table(element, doc)
                    table_data = self._extract_table_data(table)
                    tables.append(table_data)
                    
                    # Добавляем текстовое представление таблицы
                    text_parts.append("\n[ТАБЛИЦА]")
                    for row in table_data:
                        text_parts.append(" | ".join(row))
                    text_parts.append("[/ТАБЛИЦА]\n")
            
            result.text = "\n".join(text_parts)
            result.tables = tables
            result.metadata["tables_count"] = len(tables)
            
            # Проверка результата
            if len(result.text.strip()) == 0:
                result.errors.append("Документ не содержит текста")
        
        except Exception as e:
            self.logger.error(f"Ошибка парсинга DOCX: {e}")
            result.errors.append(str(e))
        
        result.parse_time = time.time() - start_time
        return result
    
    def _extract_docx_metadata(self, doc: docx.Document) -> Dict[str, Any]:
        """Извлечение метаданных DOCX"""
        metadata = {}
        try:
            core_props = doc.core_properties
            
            metadata["title"] = core_props.title or ""
            metadata["author"] = core_props.author or ""
            metadata["subject"] = core_props.subject or ""
            metadata["keywords"] = core_props.keywords or ""
            metadata["category"] = core_props.category or ""
            metadata["comments"] = core_props.comments or ""
            
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            
            metadata["paragraphs_count"] = len(doc.paragraphs)
            metadata["tables_count"] = len(doc.tables)
        
        except Exception as e:
            self.logger.warning(f"Не удалось извлечь метаданные DOCX: {e}")
        
        return metadata
    
    def _extract_table_data(self, table: Table) -> list:
        """
        Извлечение данных из таблицы
        
        Args:
            table: Объект таблицы
            
        Returns:
            Список строк таблицы
        """
        table_data = []
        try:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
        except Exception as e:
            self.logger.warning(f"Ошибка извлечения таблицы: {e}")
        
        return table_data
